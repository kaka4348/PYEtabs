def connection_export(data,tabelname,savename):
    ''' 
    data = []
    tabelname = ''
    savename = ''
    '''
    from docx import Document
    # import os
    
    # file_path = os.path.abspath(os.path.join(os.path.dirname(__file__), os.path.pardir))
    # document = Document(os.path.join(file_path, 'wordsampel', 'default.docx'))
    document = Document()

    # difine custom style
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt

    styles = document.styles
    style = styles.add_style('persian', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    style.font.bold = True
    style.font.size = Pt(12)

    styles = document.styles
    style = styles.add_style('sartitr', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    style.font.bold = True
    style.font.size = Pt(22)

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('PYEtabs', style='sartitr')

    # ejad fasele ba parageraph khali
    document.add_paragraph()
    document.add_paragraph('تهيه کننده علی فرج پورزنجانی عضويت 3556 نظام مهندسی کرمانشاه', style='persian')
    document.add_paragraph('شماره تماس   09187406759  ', style='persian')
    document.add_paragraph('آيدی تلگرام   @aliok63  ', style='persian')
    document.add_paragraph('کانال تلگرام برنامه جهت دريافت جديدترين نسخه   @PYEtabs  ', style='persian')
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    #===========================================================
    text = tabelname
    document.add_paragraph(text, style='persian')
    document.add_paragraph()
    # table
    row = len(data)
    col = len(data[-1])

    table = document.add_table(rows=row, cols=col, style ='connection')
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # from docx.enum.table import WD_TABLE_DIRECTION
    # table.direction = WD_TABLE_DIRECTION.RTL
    # from docx.enum.table import WD_ALIGN_VERTICAL
    # table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(len(data)):
        row_cells = table.rows[i].cells
        for j in range(len(data[i])):
            row_cells[j].text = str(data[i][j])

    document.save(f'{savename}.docx')

if __name__ == "__main__":
    data =[[1,2,3],
           [7,8,9],
           [10,11,12]
                    ]
    tabelname = 'لیست اتصالات خمشی:'
    savename = 'connection'

    connection_export(data,tabelname,savename)
    

