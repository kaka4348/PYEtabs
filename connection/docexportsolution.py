def connection_export_solution(data,savename):

    from docx import Document
    document = Document()

    # difine custom style
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt

    styles = document.styles
    style = styles.add_style('persian', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    style.font.bold = True
    style.font.size = Pt(12)

    styles = document.styles
    style = styles.add_style('sartitr', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    style.font.bold = True
    style.font.size = Pt(22)

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('PYEtabs', style='sartitr')

    # ejad fasele ba parageraph khali
    document.add_paragraph()
    document.add_paragraph('تهيه کننده علی فرج پورزنجانی عضويت 3556 نظام مهندسی کرمانشاه', style='persian')
    document.add_paragraph('شماره تماس   09187406759  ', style='persian')
    document.add_paragraph('آيدی تلگرام   @aliok63  ', style='persian')
    document.add_paragraph('کانال تلگرام برنامه جهت دريافت جديدترين نسخه   @PYEtabs  ', style='persian')
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    #===========================================================
    text = data
    document.add_paragraph(text, style='eng')
    document.add_paragraph()


    document.save(f'{savename}.docx')

if __name__ == "__main__":
    data =[[1,2,3],
           [7,8,9],
           [10,11,12]
                    ]
    tabelname = 'لیست اتصالات خمشی:'
    savename = 'connection'

    connection_export(data,tabelname,savename)
    

